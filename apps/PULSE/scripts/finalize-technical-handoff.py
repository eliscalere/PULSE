#!/usr/bin/env python3
"""Promote the maintained technical reference into the final developer handoff."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "02-TECHNICAL-REFERENCE.docx"
TARGET = ROOT / "docs" / "current" / "PULSE-Technical-Handoff.docx"
LOGO = ROOT / "assets" / "images" / "PULSE_logo_black_transparent_4x.png"
FONT = "Montserrat"

def fontify(run, size=None, bold=None, color=None):
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{key}"), FONT)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def clean(text):
    text = text.replace("Technical Reference", "Technical Handoff")
    text = text.replace("PULSE", "PULSE")
    text = text.replace("AEWTTR-PULSE", "PULSE")
    text = text.replace("AEWTTR ", "")
    text = text.replace("AEWTTR", "")
    text = text.replace("aewttrSaveStore", "local-mode save helper")
    return text

def update_paragraph(p):
    for run in p.runs:
        run.text = clean(run.text)
        fontify(run)

def walk(container):
    for p in container.paragraphs:
        update_paragraph(p)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                walk(cell)

def set_style(style, size, color, bold=False):
    style.font.name = FONT
    rfonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold

def style_by_name(styles, name):
    return next((style for style in styles if style.name == name), None)

def main():
    doc = Document(SOURCE)
    for section in doc.sections:
        section.top_margin = Inches(.78); section.bottom_margin = Inches(.72)
        section.left_margin = section.right_margin = Inches(.82)
        walk(section.header); walk(section.footer)
    walk(doc)
    styles = doc.styles
    normal = style_by_name(styles, "Normal")
    if normal:
        set_style(normal, 10.2, "26354A")
    for name, size, color in (("Heading 1", 17, "143F73"), ("Heading 2", 13.5, "1C5FA8"), ("Heading 3", 11.5, "1C5FA8")):
        style = style_by_name(styles, name)
        if style:
            set_style(style, size, color, True)
            style.paragraph_format.space_before = Pt(16)
            style.paragraph_format.space_after = Pt(7)
    for name in ("Title", "Subtitle", "List Bullet", "List Number", "Caption"):
        style = style_by_name(styles, name)
        if style:
            set_style(style, 11, "26354A")
    # Add the approved PULSE mark before the first title without changing the detailed body.
    logo = doc.add_paragraph(); logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo.paragraph_format.space_before = Pt(4); logo.paragraph_format.space_after = Pt(12)
    logo.add_run().add_picture(str(LOGO), width=Inches(1.35))
    body = doc._body._body
    body.remove(logo._p)
    body.insert(0, logo._p)
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET)
    print(TARGET)

if __name__ == "__main__":
    main()
