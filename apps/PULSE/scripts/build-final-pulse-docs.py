#!/usr/bin/env python3
"""Build the final PULSE documentation set with one restrained brand system."""
from pathlib import Path
import tempfile, zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "current"
LOGO = ROOT / "assets" / "images" / "PULSE_logo_black_transparent_4x.png"
FONT = "Montserrat"
NAVY, BLUE, INK, MUTED, PALE, LINE = "102A43", "246BCE", "172B4D", "5E6C84", "EEF4FB", "D7E0EA"
NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

def font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), FONT)
    for key in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
        rfonts.attrib.pop(qn(f"w:{key}"), None)
    run.font.size, run.bold, run.italic = Pt(size), bold, italic
    run.font.color.rgb = RGBColor.from_string(color)

def style_font(style, size=None, color=None, bold=None):
    style.font.name = FONT
    rf = style._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        rf.set(qn(f"w:{key}"), FONT)
    for key in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
        rf.attrib.pop(qn(f"w:{key}"), None)
    if size: style.font.size = Pt(size)
    if color: style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: style.font.bold = bold

def shade(cell, fill):
    node = OxmlElement("w:shd"); node.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(node)

def margins(cell, top=100, start=140, bottom=100, end=140):
    tcp = cell._tc.get_or_add_tcPr(); tcMar = tcp.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcp.append(tcMar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tcMar.find(qn(f"w:{side}"))
        if el is None: el = OxmlElement(f"w:{side}"); tcMar.append(el)
        el.set(qn("w:w"), str(value)); el.set(qn("w:type"), "dxa")

def borders(table, color=LINE, size="6"):
    tblPr = table._tbl.tblPr; el = tblPr.first_child_found_in("w:tblBorders")
    if el is None: el = OxmlElement("w:tblBorders"); tblPr.append(el)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{side}"); edge.set(qn("w:val"), "single"); edge.set(qn("w:sz"), size); edge.set(qn("w:color"), color); el.append(edge)

def widths(table, inches):
    table.autofit = False
    grid = table._tbl.tblGrid
    for idx, width in enumerate(inches):
        grid.gridCol_lst[idx].set(qn("w:w"), str(int(width * 1440)))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(inches[idx])
            tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW(); tcW.set(qn("w:w"), str(int(inches[idx] * 1440))); tcW.set(qn("w:type"), "dxa")

def configure(doc, name):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(.78); sec.left_margin = sec.right_margin = Inches(.82)
    sec.header_distance = sec.footer_distance = Inches(.32)
    normal = doc.styles["Normal"]; style_font(normal, 10.3, INK); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.18
    for key in ("Title", "Subtitle", "Caption", "List Bullet", "List Number"): style_font(doc.styles[key])
    for key, size, col, before, after in (("Heading 1", 17, NAVY, 18, 8), ("Heading 2", 12.5, BLUE, 12, 5), ("Heading 3", 11, NAVY, 8, 3)):
        st = doc.styles[key]; style_font(st, size, col, True); st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True
    h = sec.header.paragraphs[0]; h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run("PULSE  /  " + name.upper()); font(r, 8.1, True, MUTED)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("PULSE  •  INTERNAL  •  "); font(r, 8, False, MUTED)
    field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); footer._p.append(field)

def p(doc, text="", lead=None, size=10.3, after=None):
    out = doc.add_paragraph();
    if after is not None: out.paragraph_format.space_after = Pt(after)
    if lead and text.startswith(lead):
        font(out.add_run(lead), size, True); font(out.add_run(text[len(lead):]), size)
    else: font(out.add_run(text), size)
    return out

def heading(doc, text, level=1):
    h = doc.add_heading(text, level)
    for r in h.runs: font(r, {1:17,2:12.5,3:11}[level], True, {1:NAVY,2:BLUE,3:NAVY}[level])
    return h

def image(doc, file, caption):
    path = OUT / "screenshots" / file
    if path.exists():
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_before = Pt(6); q.paragraph_format.space_after = Pt(2)
        q.add_run().add_picture(str(path), width=Inches(6.25))
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_after = Pt(8)
        font(q.add_run(caption), 8.5, False, MUTED)

def bullets(doc, items):
    for item in items:
        q = doc.add_paragraph(style="List Bullet"); q.paragraph_format.space_after = Pt(3); q.paragraph_format.line_spacing = 1.12; font(q.add_run(item), 10.1)

def steps(doc, items):
    for item in items:
        q = doc.add_paragraph(style="List Number"); q.paragraph_format.space_after = Pt(4); q.paragraph_format.line_spacing = 1.12; font(q.add_run(item), 10.1)

def label_box(doc, label, text, fill=PALE):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; widths(t, [6.82]); borders(t, fill, "4")
    c = t.cell(0,0); shade(c, fill); margins(c, 130, 180, 130, 180)
    q = c.paragraphs[0]; q.paragraph_format.space_after = Pt(0); font(q.add_run(label.upper() + "  "), 8.1, True, BLUE); font(q.add_run(text), 9.8, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def table(doc, headers, rows, w):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; widths(t,w); borders(t)
    for idx, label in enumerate(headers):
        c=t.rows[0].cells[idx]; shade(c,NAVY); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(0); font(q.add_run(label),8.4,True,"FFFFFF")
    for values in rows:
        cells=t.add_row().cells
        for idx,val in enumerate(values):
            c=cells[idx]; margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(0); font(q.add_run(str(val)),8.9)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def cover(doc, title, subtitle, audience, purpose):
    q=doc.add_paragraph(); q.paragraph_format.space_before=Pt(38); q.alignment=WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists(): q.add_run().add_picture(str(LOGO), width=Inches(1.7))
    q=doc.add_paragraph(); q.paragraph_format.space_before=Pt(30); q.paragraph_format.space_after=Pt(8); font(q.add_run("DOCUMENTATION SERIES"),8.3,True,BLUE)
    q=doc.add_paragraph(); q.paragraph_format.space_after=Pt(10); font(q.add_run(title),28,True,NAVY)
    q=doc.add_paragraph(); q.paragraph_format.space_after=Pt(22); font(q.add_run(subtitle),13,False,MUTED)
    label_box(doc,"Purpose", purpose)
    t=doc.add_table(rows=2, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(t,[3.41,3.41]); borders(t,"FFFFFF","0")
    values=(("AUDIENCE",audience),("STATUS","Final delivery"),("VERSION","2026.07"),("REVIEW","After a material change"))
    for idx,(k,v) in enumerate(values):
        c=t.cell(idx//2,idx%2); shade(c,"F7F9FC"); margins(c,150,160,150,160)
        q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(0); font(q.add_run(k+"\n"),7.6,True,MUTED); font(q.add_run(v),9.3,True,INK)
    doc.add_page_break()

def flow(doc, title, items):
    heading(doc,title,2); t=doc.add_table(rows=1,cols=len(items)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(t,[6.82/len(items)]*len(items)); borders(t,"FFFFFF","0")
    for idx,item in enumerate(items):
        c=t.cell(0,idx); shade(c, PALE if idx%2==0 else "E2ECF9"); margins(c,145,110,145,110)
        q=c.paragraphs[0]; q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_after=Pt(0); font(q.add_run(item),8.4,True,NAVY)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def user_doc():
    d=Document(); configure(d,"User Guide")
    cover(d,"User Guide","The practical path through daily work, project delivery, meetings, travel, and reviews.","Authorized users","Use this guide as a quick working reference. It explains where to act, what creates an official record, and when to request help.")
    heading(d,"1. Begin each work session")
    image(d,"01-dashboard-current.png","Dashboard — Start-of-session attention items.")
    p(d,"Open PULSE from the approved SharePoint page. Your existing SharePoint session establishes your access; there is no separate PULSE password.")
    flow(d,"Your opening check",["Confirm name\nand role","Review attention\nitems","Open My Work\nand priorities","Act, assign,\nor schedule"])
    bullets(d,["Use Dashboard as the start-of-session checkpoint.","Use Overview → My for assigned work and Overview → Team for the portfolio-level summary.","Use the bell and user menu for attention items and notification preferences.","If the application opens outside SharePoint, local fallback is non-persistent. Do not treat that data as an official record."])
    heading(d,"2. Find the right workspace")
    image(d,"02-overview-current.png","Overview — Live portfolio summary.")
    table(d,["Need","Use","What to do"],[("Portfolio health","Overview","Read the summary; make edits in the underlying workspace."),("Project delivery","Projects","Maintain scope, roles, dates, tracker work, risks, and project context."),("Meeting decisions","Weekly Meeting","Capture decisions and convert follow-ups into owned actions."),("Travel request","Travel","Create, submit, monitor, and update the request."),("Formal concurrence","Document Review","Route the correct revision to named reviewers and preserve history.")],[1.35,1.45,4.02])
    heading(d,"3. Maintain project work")
    image(d,"03-project-tracker-current.png","Project Tracker — Record actionable dates and health.")
    steps(d,["Open Projects and select the required project.","Confirm project health, dates, lifecycle details, and named roles before changing records.","Create or update tracker work with an owner, practical date, status, and clear execution context.","Use Blocked only when work cannot proceed. State the dependency and raise the recovery action in the appropriate meeting."])
    label_box(d,"Working rule","Boards and checklists help visualize a workflow. The tracker is the authoritative work list for ownership, dates, and portfolio health.")
    heading(d,"4. Run a meeting that creates follow-through")
    image(d,"04-weekly-meeting-current.png","Weekly Meeting — Decisions and carry-forward actions.")
    steps(d,["Confirm the meeting scope and participant roster.","Review prior actions, priorities, and blocked work.","Capture material decisions and discussion notes.","Create or update named tracker actions before closing the meeting.","Verify each action has an owner and date in the appropriate tracker."])
    heading(d,"5. Submit travel and route reviews")
    image(d,"05-travel-request-current.png","Travel Request — Enter required details.")
    table(d,["Workflow","Before submission","After submission"],[("Travel","Enter accurate purpose, dates, travelers, location, estimates, and support detail.","Monitor status, update material changes, and complete any required debrief."),("Document Review","Use the current revision, due date, context, and required reviewers.","Monitor decisions and feedback; preserve the revision trail before advancing.")],[1.25,2.78,2.79])
    image(d,"06-document-review-current.png","Document Review — Track formal concurrence.")
    label_box(d,"Important","Calendar visibility is not authorization. A file stored in a project library is not a formal concurrence record.","FFF5E5")
    heading(d,"6. Get help safely")
    p(d,"For a defect, capture the time, SharePoint page or route, affected record identifier, your role, expected result, actual result, and an approved screenshot. Refresh first and confirm you are in the approved SharePoint-hosted version. Do not alter roles or schema unless you are authorized to administer them.")
    save(d,OUT/"PULSE-User-Guide.docx")
    write_md("PULSE-User-Guide.md",USER_MD)

def sop_doc():
    d=Document(); configure(d,"Operations SOP")
    cover(d,"Operations SOP","Detailed, in-app procedures for creating, routing, updating, approving, and closing operational records.","Users, project leads, travelers, reviewers, approvers, and administrators","Use this document at the point of work. It identifies the PULSE workspace, the fields that make each record usable, the allowed handoffs, and the record that must exist at completion.")
    heading(d,"How to use this SOP")
    p(d,"A workflow starts in the named workspace and ends only when the stated record is saved and visible in PULSE. Required below means required for a usable operational record even where the screen permits a temporary draft or blank value.")
    table(d,["Term","Meaning"],[("Record owner","Person accountable for accurate fields and follow-through."),("Official record","The saved PULSE record in its workspace; supporting files alone do not replace it."),("Status","The current workflow state shown to other users; update it when the work state changes."),("Control","The check that prevents an incomplete, unauditable, or misleading record.")],[1.45,5.37])
    heading(d,"SOP 01 — Create and maintain a project")
    p(d,"Workspace: Projects  |  Owner: project lead or delegate  |  Use: when a new project, subproject, or portfolio workstream needs an accountable home.",size=9.3)
    heading(d,"Create the project",2)
    steps(d,["Navigate to Projects and select New Project.","Enter Project name (reader-facing), Portfolios, and a clear Description.","Add Members using the people picker to establish the initial roster. (Optionally add a cover image).","Save the new project, then open it to configure Settings: set Lifecycle status, Technical health, and essential Dates (start, due, completion) before assigning work."])
    table(d,["Project field","Enter","Use it for"],[("Project name","Clear project or effort name","Navigation, reports, and the project banner."),("Portfolios","Applicable portfolio names","Portfolio rollups and filtering."),("Description","Plain-language purpose and expected outcome","Shared context for users entering the project."),("Members","People or groups who participate","Roster, meeting participation, and owner selection."),("Settings → status","Lifecycle and current technical health","Portfolio health and escalation views."),("Settings → dates","Start, due, and completion dates","Schedule views and overdue identification."),("Settings → roles","Named project roles and contacts","Accountability and handoff context.")],[1.6,2.65,2.57])
    label_box(d,"Control","Before creating tracker work, verify the project name, portfolio, members, lifecycle status, and due date. Do not use a project description as a substitute for an owner or a dated task.")
    heading(d,"Maintain tracker work and risks",2)
    image(d, "03-project-tracker-current.png", "Current Project Tracker.")
    steps(d,["Within the project, open the Tracker and create or edit the work item.","Give the item a concise action title, then assign an owner and define Start date, Due date, Action status, and Health.","Use Not Started, In Progress, Blocked, On Hold, Complete, Done, or Cancelled for Action status.","Use Health to communicate On Track, At Risk, Off Track, Blocked, On Hold, or Complete.","If a task is Blocked, you must enter the dependency in the Blocked reason field. If On Hold, enter the On-hold reason.","For a risk, enter a name, description, owner, likelihood, impact, category, mitigation plan, due date, response strategy, and current risk status.","Review red or amber project health, blocked work, overdue dates, and open risks during the project meeting; update the record instead of leaving the change only in notes."])
    label_box(d,"Completion","Project information is current; every active task has an owner, status, and usable date; every blocked or high-risk condition has a documented cause and recovery path.","F7F9FC")
    heading(d,"SOP 02 — Run a weekly or project meeting")
    p(d,"Workspace: Weekly Meeting or Project → Meeting  |  Owner: facilitator or project lead  |  Use: scheduled operational review and decision capture.",size=9.3)
    image(d, "04-weekly-meeting-current.png", "Weekly Meeting.")
    steps(d,["Open the correct meeting scope. Use Weekly Meeting for the general series; use the project Meeting tab when the discussion is project-specific.","Set or confirm Meeting date, Meeting series, and attendees. Use the project roster as the default participant source for a project meeting.","Review the agenda, current Rocks, tracker items, risks, and carry-forward work before beginning discussion.","During the meeting, enter Notes for context; record Decisions separately when a choice, approval, or direction must be traceable.","For every follow-up, create or update the underlying tracker item with owner, due date, status, health, and notes. Do not rely only on Action items in meeting notes.","Record risk changes and project status changes when the meeting changes a condition visible to the portfolio.","End the session only after the follow-ups are visible in the tracker and the meeting Session status is ended."])
    table(d,["Meeting field","What it records","Rule"],[("Meeting date / series","Which session the record represents","Use the scheduled date and a recognizable series name."),("Attendees / attendance","Who was invited or present","Keep the roster aligned to actual participants."),("Agenda","Topics to be covered","Use it to structure the session, not to store decisions."),("Notes","Discussion context","Summarize facts, issues, and rationale."),("Decisions","Traceable choices or approvals","State the decision and owner/context."),("Action items / carry-forward","Follow-up work","Create a corresponding tracker item when action is required."),("Session status","planned, active, or ended","End only when notes and actions are complete.")],[1.65,2.52,2.65])
    label_box(d,"Control","Meeting notes do not replace owned, dated tracker work. A decision without an owner or a resulting record is not closed.")
    heading(d,"SOP 03 — Submit and manage travel")
    p(d,"Workspace: Travel → New Request  |  Owner: requester; approver and finance roles act at the assigned stage  |  Use: standard, engineering, project, contractor, or leave requests.",size=9.3)
    heading(d,"Create the request",2)
    image(d, "05-travel-request-current.png", "Travel Request Wizard.")
    steps(d,["Choose the correct request type. Select the specialized engineering form only when the request requires its additional travel details.","Enter the trip title, traveler roster, destination, start date, end date, time information when applicable, purpose, and impact if not approved.","Select related project records when travel supports project work. Enter travel type, estimate, alternatives considered, and travel notes when those fields apply.","For an engineering request, complete its additional traveler, location, date, transportation, and funding detail fields before submission.","Review the summary, then submit. The system assigns the request identifier and creates the Travel request record.","Use My Travel to monitor a request you submitted or are traveling on. While it is pending, correct it through Edit or withdraw it through Revoke when appropriate."])
    table(d,["Travel field","What to enter","Why it matters"],[("Trip title / Request type","Clear purpose and correct form category","Identifies the request and drives the applicable form."),("Travelers","Requester plus every traveler","Controls visibility and links individual debriefs."),("Destination / dates / times","Actual location and planned travel window","Calendar, duration, and approval context."),("Purpose / impact","Business need and consequence of non-approval","Decision rationale for the approver."),("Project / travel type","Related project and TDY, conference, training, or other","Reporting and operational context."),("Estimate / alternatives / notes","Cost context, options, and restrictions","Supports an informed approval."),("Charge object","Finance assignment when required","Completes the finance handoff.")],[1.55,2.9,2.37])
    heading(d,"Approve, assign finance, or close",2)
    steps(d,["Approvers use Travel → All Travel. Review the identifier, travelers, destination, dates, purpose, impact, estimate, project, and attached form details before acting.","For a non-leave request in Pending, select Approve to move it to Pending Finance, or Deny and enter a clear denial reason. Leave requests follow their concurrence action.","Finance users open Pending Finance, enter the charge object, and finalize the approval. The request then shows Approved.","Requesters may cancel an approved request they no longer need; authorized administrators may cancel when required. Do not overwrite history to represent cancellation.","After approved travel, each traveler opens Debrief, selects the request, and enters trip dates, systems or subjects, classification, summary, and follow-up. Save one debrief per traveler."])
    label_box(d,"Status path","Pending → Pending Finance → Approved. Other terminal or corrective states are Denied, Revoked, and Cancelled. Calendar visibility is not authorization evidence.","FFF5E5")
    heading(d,"SOP 04 — Route and complete document review")
    p(d,"Workspace: Document Review  |  Owner: document owner; reviewers and signers perform assigned actions  |  Use: auditable review, concurrence, final-pack, and signature workflows.",size=9.3)
    image(d, "06-document-review-current.png", "Document Review Board.")
    steps(d,["Navigate to Document Review and click New Review. Select the related project, document type, and kind. Provide the Title, Due Date, and Revision Label (e.g. v1.0).","Upload the Current File representing the exact version to be reviewed. The system retains revision history.","Add individual Reviewers or a reviewer group. Designate specific reviewers as Signers only if sequential signatures are required after review.","Submit the review. Reviewers will receive notifications and must log their decision (Approved/Changes Requested) along with any actionable comments against the current revision.","If any reviewer requests changes, the owner must update the local file, upload it as a new revision, and restart the review cycle for the new revision. Do not overwrite history.","When all reviewers approve, the record enters Review Complete when no signature is required.","If signatures are required, upload the Final Pack PDF to move it to Awaiting Final Pack and begin sequential signing.","Signers download the pack, sign externally, and upload the signed file to advance the record to Signed."])
    table(d,["Document-review field","What to enter","Control"],[("Project / document type / kind","Business context and record class","Use the correct project and document classification."),("Owner / submitter / deadline","Accountability, source, and due date","Owner is accountable for routing and follow-up."),("Revision label / current file","The exact revision under review","New revision resets review decisions for the new content."),("Reviewers / reviewer group","People required to decide","Use named people; group selection expands to people."),("Signers / signing order","Sequential signatory roster","Lock and preserve the order once signing begins."),("Comments / decisions","Actionable review outcome","Comments must explain requested change or approval rationale."),("Final pack","PDF used for signature stage","Upload only after review completion; it is not a draft revision.")],[1.65,2.45,2.72])
    label_box(d,"Status path","Not Started → In Review → Changes Requested or Review Complete. Signature-required records continue through Awaiting Final Pack, Signing in Progress, and Signed; archive only after the record is complete.","FFF5E5")
    heading(d,"SOP 05 — Administer access and setup")
    p(d,"Workspace: Admin  |  Owner: authorized administrator  |  Use: onboarding, offboarding, role maintenance, and controlled setup.",size=9.3)
    steps(d,["Open Admin and confirm current-user and target-site diagnostics before changing roles or schema.","In Users, create or update the person’s PULSE App Role. Maintain user email, display name, active state, role, job title, access selections, and notification preferences as applicable.","Choose the least-privileged role that supports the person’s work: Admin, Meeting Admin, Finance Admin, Document Admin, Manager, Member, or Viewer.","For offboarding, set the role record inactive rather than deleting the operational history associated with the user.","Use SharePoint Setup to validate or provision the required lists and fields. Do not manually improvise a list or column as a substitute for the schema.","Review logs or activity records after sensitive changes and retain at least one active recovery-capable administrator."])
    label_box(d,"Control","Site membership alone is not the normal application authorization model. Role, active state, and the live role record must agree before an access issue is considered resolved.")
    heading(d,"Incident reporting minimum")
    p(d,"For any workflow defect, record the time, workspace and route, affected project/request/document identifier, user role, exact field or action, expected result, actual result, browser or device, and approved screenshot. Preserve logs or activity evidence. Do not repair production roles or schema while the cause is still unknown.")
    save(d,OUT/"PULSE-Operations-SOP.docx")
    write_md("PULSE-Operations-SOP.md",SOP_MD_DETAILED)

def tech_doc():
    d=Document(); configure(d,"Technical Handoff")
    cover(d,"Technical Handoff","Developer continuity guide: runtime, source layout, SharePoint data, packaging, deployment, validation, and recovery.","Developers, site administrators, and release managers","This is the developer pickup document. Read it before making a code or schema change; it describes the actual boundaries that keep the application deployable.")
    heading(d,"1. Operating model")
    p(d,"PULSE is a browser-only single-page application. It is hosted as one self-contained HTML file in SharePoint and runs inside the visitor’s existing SharePoint session. SharePoint Lists and document libraries provide the shared operational record. There is no application server, separate database, separate login, or required runtime CDN.")
    flow(d,"Runtime path",["Browser +\nSharePoint session","One-file PULSE\nHTML package","Shell + feature\nmodules","REST adapter +\nrepository","SharePoint lists\nand libraries"])
    label_box(d,"Non-negotiable boundary","Keep application data on SharePoint. Do not introduce a required backend, Azure registration, Graph dependency, external API, or runtime CDN into the shipping artifact.","FFF5E5")
    heading(d,"2. Repository map")
    table(d,["Area","Paths","Responsibility"],[("Entry","index.html","Development entry point; determines CSS, vendor, and script load order."),("Application shell","assets/js/app.js; data.js","Boot, shared state, routing, rendering shell, in-memory data shape."),("Feature modules","assets/js/pages/","One module per user-facing workspace and its interactions."),("SharePoint layer","sharepoint-adapter.js; sharepoint-repo.js; sharepoint-schema.js","REST plumbing, object/list mapping, save semantics, schema/provisioning."),("Packaging","scripts/build-sharepoint-package.js","Preferred source-to-one-file ship artifact."),("Alternative wrapper","scripts/build-forge.js","Secondary portable/iframe wrapper; not the ordinary hosted release artifact.")],[1.25,2.55,3.02])
    heading(d,"3. Source loading and application boot")
    p(d,"The source entry loads vendor libraries first, then configuration and SharePoint code, then data/services, then the shell and page modules. The order is material because the source uses browser globals rather than a module bundler. Do not reorder scripts casually.")
    bullets(d,["app-config.js supplies site overrides, list prefix, and boot diagnostics.","sharepoint-adapter.js owns raw SharePoint REST calls, site detection, current-user lookup, request digest handling, and diagnostics.","sharepoint-schema.js centralizes lists, columns, validation, and setup/provisioning.","sharepoint-repo.js is the boundary page modules use to save and remove records; page code should not make raw REST calls.","data.js normalizes the in-memory store; app.js starts boot, maintains shared state, resolves hash routes, renders the shell, and dispatches page renderers."])
    heading(d,"4. SharePoint integration")
    p(d,"The adapter uses same-origin fetch calls to /_api/web/... with the visitor’s SharePoint session. Site detection favors the injected SharePoint context, then the configured manual site URL, then a limited origin-based fallback. When no site is detected, the application uses local fallback for UI exercise only; it is not an official shared record.")
    table(d,["Responsibility","Where","What to preserve"],[("Identity","Adapter + boot","Current-user REST identity is merged with page context; do not create a second authentication path."),("Roles","PULSE App Roles list","Resolve app roles from the role record; ordinary authorization is not inferred from site membership."),("Schema","SHAREPOINT_SCHEMA","Add lists/columns here first; validate/provision through setup instead of ad hoc manual fields."),("Persistence","Repo.save / Repo.remove","Keep object-to-item mappers and item-to-object mappers symmetric, especially JSON fields."),("Errors","Repository diagnostics","Persist the last SharePoint failure for recovery screens and capture the live request before changing UI.")],[1.45,1.8,3.57])
    heading(d,"5. Data model and save behavior")
    p(d,"The application loads list records into an in-memory store. Feature pages operate on those objects and call the repository when a record changes. The repository maps an object to one list item, performs the REST operation, then refreshes or reconciles the local state.")
    flow(d,"Write path",["Feature page\nchanges object","Repository maps\nobject to fields","Adapter issues\nsame-origin REST","List item saves\nthen state refreshes"])
    bullets(d,["Project IDs are maintained as a project-code ↔ SharePoint-item map; preserve this when changing project joins.","Several compound values live in *Json fields. Treat their shape as a compatibility contract: read, write, and live validation must change together.","Writes are debounced and chained per object; the effective policy is last-write-wins. Avoid parallel UI changes to the same record without deliberate coordination.","If a record is wrong, inspect the list item, mapper, JSON shape, and cached object before changing the page renderer."])
    heading(d,"6. Safe code-change workflow")
    steps(d,["Locate the page module, related store shape, repository mapper, and schema definition before editing.","Make the smallest coherent source change. A new field normally requires schema, list mapping, object normalization, UI validation, and live read/write testing.","Run the unbundled source locally for UI behavior. Local fallback is useful for layout and interaction, not for permission or persistence validation.","Use a safe SharePoint site to test identity, role resolution, schema validation, list access, and the real persistence path.","Build a one-file package, upload it, and test the actual hosted page.","Record the package name, validation evidence, and previous known-good artifact before declaring release complete."])
    heading(d,"7. Build the shipping file")
    p(d,"The preferred builder is scripts/build-sharepoint-package.js. It converts the multi-file source tree into the flat HTML file expected by the hosted page.")
    label_box(d,"Command","node scripts/build-sharepoint-package.js \"FS packages/PULSE_v.<YYYY.MM.DD>.html\"","EAF2FC")
    table(d,["Builder action","Why it matters"],[("Reads index.html","Preserves the authoritative development dependency order."),("Inlines styles and rewrites CSS URLs","The released file carries its CSS, fonts, and image dependencies."),("Inlines scripts and escapes closing script text","Avoids broken inline script blocks in the generated HTML."),("Strips local cache-busting query strings","Query strings on hosted asset paths can produce silent failures; none must remain in the ship file."),("Embeds selected images as data URIs","The package remains self-contained."),("Wraps the presentation library UMD bundle","Prevents an AMD host from intercepting the browser-global registration."),("Writes a provenance manifest comment","Hashes and generation time provide release traceability."),("Fails external styles/scripts","A release must not depend on an external runtime resource.")],[2.85,4.0])
    heading(d,"8. Builder choice and artifact discipline")
    p(d,"Do not substitute the secondary wrapper builder for an ordinary hosted release. The secondary script creates a heavier iframe/srcdoc-style wrapper for portable distribution. The preferred builder produces the direct flat package used for the normal SharePoint-hosted deployment.")
    bullets(d,["Package after every source change; there is no automatic deploy or watch path.","Use a dated package name and retain the prior known-good package as a rollback candidate.","The embedded manifest generation time is more reliable than a filename alone when confirming provenance.","Never upload an untested local artifact as the only recovery option."])
    heading(d,"9. Deployment and acceptance")
    steps(d,["Build the package into FS packages with a clear dated name.","Upload the generated HTML to the document library used by the hosted page.","Open the real hosted page—not a local tab—and verify initial load, current user, role, navigation, and the affected workflow.","Validate a real read and write for the changed list or feature. Confirm expected list fields after save.","Check the browser console and the application diagnostics for errors.","Retain validation evidence and the prior package. If validation fails, restore the prior package first, then investigate."])
    heading(d,"10. Troubleshooting playbook")
    table(d,["Symptom","First inspection","Safe response"],[("Blank or partial app","Hosted page, console, generated package, asset query strings","Rebuild with the preferred builder; verify the package has no external runtime dependency."),("Wrong role or no admin access","Current-user REST, detected site, role list record","Fix the active role record or site configuration; do not assume site membership grants the app role."),("Save fails","Logs, permissions, list schema, mapper, request payload","Validate the live REST response and field mapping before changing the screen."),("Data appears wrong","List item, JSON shape, mapper pair, local cache","Correct the data boundary coherently and run read/write validation."),("Release-only defect","Actual hosted page and prior package","Rollback to the known-good package, capture evidence, then isolate the source change.")],[1.3,2.6,2.95])
    heading(d,"11. Developer handoff checklist")
    bullets(d,["Understand the one-file hosted runtime and its no-backend/no-CDN constraints.","Trace each changed feature from page module to store, mapper, schema, and live list.","Use the preferred package builder for the normal release path.","Validate in a safe SharePoint environment and then in the actual hosted page.","Keep a known-good package and record release evidence for every production change."])
    save(d,OUT/"PULSE-Technical-Handoff.docx")
    write_md("PULSE-Technical-Handoff.md",TECH_MD)

def save(doc,path):
    doc.save(path)
    with zipfile.ZipFile(path,"r") as source: content={x.filename:source.read(x.filename) for x in source.infolist()}
    root=ET.fromstring(content["word/theme/theme1.xml"])
    for item in root.findall(f".//{NS}majorFont")+root.findall(f".//{NS}minorFont"):
        for tag in ("latin","ea","cs"):
            el=item.find(f"{NS}{tag}")
            if el is not None: el.set("typeface",FONT)
    content["word/theme/theme1.xml"]=ET.tostring(root,encoding="utf-8",xml_declaration=True)
    with tempfile.NamedTemporaryFile(dir=path.parent,suffix=".docx",delete=False) as fh: tmp=Path(fh.name)
    with zipfile.ZipFile(tmp,"w",zipfile.ZIP_DEFLATED) as target:
        for name,data in content.items(): target.writestr(name,data)
    tmp.replace(path)

def write_md(name,text): (OUT/name).write_text(text.strip()+"\n",encoding="utf-8")

USER_MD='''# PULSE User Guide

## Start each session

Open PULSE from the approved SharePoint page. Your existing SharePoint session establishes access; there is no separate password. Review Dashboard, attention items, My Work, and red or amber projects. Local fallback outside SharePoint is non-persistent and not an official record.

## Find the right workspace

| Need | Workspace | Action |
|---|---|---|
| Portfolio health | Overview | Read the summary; edit in the underlying workspace. |
| Project delivery | Projects | Maintain roles, dates, tracker work, risks, and context. |
| Meeting follow-through | Weekly Meeting | Capture decisions and create owned actions. |
| Travel | Travel | Create, submit, update, and monitor the request. |
| Formal concurrence | Document Review | Route the correct revision to named reviewers. |

## Core rules

1. Use the tracker for authoritative ownership, dates, and health.
2. Mark work Blocked only when a dependency prevents progress; name the dependency and recovery action.
3. Convert meeting follow-ups into named tracker actions before closing.
4. Calendar visibility is not authorization; project-library storage is not formal concurrence.
5. Preserve document revision history.

## Support

Report time, SharePoint page or route, record identifier, your role, expected result, actual result, and an approved screenshot. Do not alter roles or schema unless authorized.
'''

SOP_MD='''# PULSE Operations SOP

## Control procedures

1. **Start each session:** confirm identity and role; review attention items and active work.
2. **Maintain projects:** keep roles, tracker ownership, dates, status, causes, and recovery actions current.
3. **Conduct meetings:** capture material decisions and create named dated actions before close.
4. **Travel:** submit complete business details; approvers decide with an explanatory note; update material changes.
5. **Document Review:** route the current revision to required reviewers and preserve the review history.
6. **Administration:** manage app roles through the approved model and validate/provision schema through setup.
7. **Release:** test locally and safely in SharePoint, build one file, validate the real hosted page, and retain rollback.

## Incident reporting

Capture time, page, user role, route, record identifier, expected and actual result, browser/device, and approved screenshot. Preserve logs and audit evidence. Do not make uncoordinated production role or schema changes while troubleshooting.
'''

SOP_MD_DETAILED='''# PULSE Operations SOP

## 1. Project workflow

![Project Tracker](screenshots/03-project-tracker-current.png)

1. Navigate to **Projects** and click **New Project**.
2. Enter **Project name** (reader-facing), **Portfolios**, and a concise **Description**.
3. Add **Members** via the people picker to establish the initial roster. (Optionally add a cover image).
4. Save the project, then open it to configure **Settings**: set the **Lifecycle status**, **Technical health**, and essential **Dates** (start, due, completion).
5. For tracker work, assign each task to an owner and define the **Start date**, **Due date**, **Action status**, and **Health**.
6. If a task is Blocked, you must enter the dependency in the **Blocked reason**. If On Hold, enter the **On-hold reason**.
7. Create **Risk** entries with a defined likelihood, impact, category, mitigation plan, and owner.

A project is operationally current only when active work has owners and dates and red/amber conditions have a documented recovery action.

## 2. Meeting workflow

![Weekly Meeting](screenshots/04-weekly-meeting-current.png)

Use **Weekly Meeting** for the general series or **Project → Meeting** for a project-scoped session. Maintain meeting date, series, attendees, agenda, notes, decisions, action items, carry-forward work, risk changes, project status changes, and session status.

1. Confirm session scope, date, series, and roster.
2. Review rocks, tracker work, risks, and carry-forward items.
3. Record decisions separately from general notes.
4. Create or update a tracker item for every actionable follow-up.
5. End the session only after actions are owned and dated in the tracker.

## 3. Travel workflow

![Travel Request](screenshots/05-travel-request-current.png)

Start at **Travel → New Request**. Select the right request type, then enter trip title, travelers, destination, dates/times, purpose, impact if not approved, related project, travel type, estimate, alternatives, and notes. Engineering travel also requires its additional location, date, traveler, transportation, and funding detail.

The requester submits and monitors in **My Travel**. Status is **Pending**, then **Pending Finance** after approver approval, then **Approved** after finance records the charge object. An approver can deny with a denial reason; pending requests may be revoked, and approved requests may be cancelled. Each traveler completes a separate debrief with trip dates, systems/subjects, classification, summary, and follow-up.

## 4. Document Review workflow

![Document Review](screenshots/06-document-review-current.png)

1. Navigate to **Document Review** and click **New Review**.
2. Select the related **Project**, document type, and kind. Provide the **Title**, **Due Date**, and **Revision Label** (e.g. v1.0).
3. Upload the **Current File** representing the version to be reviewed.
4. Add **Reviewers** (individuals or groups). If sequential signatures are required, designate specific reviewers as Signers.
5. Reviewers will receive notifications and must log their decision (Approved/Changes Requested) along with any actionable comments.
6. If changes are requested, the owner must update the local file, upload it as a new revision, and restart the review cycle for the new revision. Do not overwrite the approved revision.
7. Once all reviewers approve, the record enters **Review Complete**. If signatures are required, upload the **Final Pack** PDF to move it to **Awaiting Final Pack** and begin sequential signing.
8. Signers download the pack, sign externally, and upload the signed file to advance the record to **Signed**.

## 5. Administration workflow

![Admin](screenshots/07-admin-current.png)

In **Admin**, first confirm current-user and site diagnostics. Maintain active role records with user email, display name, role, job title, access selections, and notification preferences. Use the least-privileged applicable role. Set a departing user inactive rather than removing the record. Use **SharePoint Setup** to validate/provision required lists and fields; do not create ad hoc schema changes.

## Incident report

Capture time, workspace/route, record identifier, user role, field/action, expected result, actual result, browser/device, screenshot, and relevant log evidence.
'''

TECH_MD='''# PULSE Technical Handoff

## Operating model

PULSE is a browser-only single-page application hosted as a self-contained HTML file in SharePoint. The existing SharePoint session supplies identity; SharePoint Lists and libraries supply shared records. There is no application server, separate database, separate login, required runtime CDN, or required external API.

```text
Browser + SharePoint session → one-file PULSE package → shell + feature modules
→ SharePoint REST adapter + repository → SharePoint lists and libraries
```

## Source structure

| Area | Primary paths | Responsibility |
|---|---|---|
| Entry | `index.html` | Development dependency order. |
| Shell/store | `assets/js/app.js`, `data.js` | Boot, routing, rendering shell, shared state. |
| Features | `assets/js/pages/` | Workspace-specific UI and interactions. |
| SharePoint | `sharepoint-adapter.js`, `sharepoint-repo.js`, `sharepoint-schema.js` | REST, mapping, persistence, schema. |
| Preferred packaging | `scripts/build-sharepoint-package.js` | One-file release artifact. |

## Change workflow

1. Trace page module → store → repository mapper → schema before editing.
2. Change schema, object/list mappers, normalization, UI, and live validation together when adding data.
3. Test UI locally; use a safe SharePoint site for identity, role, schema, and persistence checks.
4. Build the one-file package, then validate in the actual hosted page.
5. Retain the prior known-good package and release evidence.

## Packaging

Use:

```bash
node scripts/build-sharepoint-package.js "FS packages/PULSE_v.<YYYY.MM.DD>.html"
```

The preferred builder reads `index.html`, inlines styles and scripts, rewrites CSS URLs and selected images to data URIs, strips development cache queries, protects the presentation-library UMD global, writes a provenance manifest, and fails external runtime dependencies. The secondary wrapper builder is for portable/iframe distribution, not the normal hosted release.

## Persistence and troubleshooting

Feature modules change in-memory objects and call the repository. The repository maps objects to list items, sends same-origin REST calls, and refreshes/reconciles local state. JSON fields are a compatibility contract; change both mapper directions and validate a live read/write. Saves are debounced and chained per object, with effective last-write-wins behavior.

For blank pages, inspect the real hosted page, console, package, and asset queries. For save failures, inspect the live REST response, permissions, list schema, mapper, and payload before changing UI. For a release-only defect, roll back to the prior package, capture evidence, and isolate the source change.
'''

if __name__ == "__main__":
    OUT.mkdir(parents=True,exist_ok=True)
    user_doc(); sop_doc(); tech_doc()
    print("Built final PULSE documentation set.")
