#!/usr/bin/env python3
"""Build the current PULSE documentation Word set from maintained content."""
from pathlib import Path
import tempfile
import zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "current"
SHOT = OUT / "screenshots"
DIAGRAM = ROOT / "docs" / "diagrams"
LOGO = ROOT / "assets" / "images" / "PULSE_logo_black_transparent_4x.png"
FONT = "Montserrat"
BLUE = "2E5AAC"
INK = "1F2937"
MUTED = "5B677A"
PALE = "EAF0FB"
DRAWINGML = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

def apply_style_font(style, family=FONT):
    """Apply the document typeface across every script supported by Word."""
    style.font.name = family
    rfonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for script in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{script}"), family)
    for theme in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
        rfonts.attrib.pop(qn(f"w:{theme}"), None)

def set_font(run, size=10.2, bold=False, color=INK, family=FONT):
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:ascii"), family)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), family)
    run._element.rPr.rFonts.set(qn("w:cs"), family)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    for theme in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
        run._element.rPr.rFonts.attrib.pop(qn(f"w:{theme}"), None)
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)

def save_doc(doc, path):
    """Save a DOCX with Montserrat as both explicit and theme typography."""
    doc.save(path)
    with zipfile.ZipFile(path, "r") as source:
        content = {item.filename: source.read(item.filename) for item in source.infolist()}
    theme_name = "word/theme/theme1.xml"
    root = ET.fromstring(content[theme_name])
    for font_set in root.findall(f".//{DRAWINGML}majorFont") + root.findall(f".//{DRAWINGML}minorFont"):
        for script in ("latin", "ea", "cs"):
            element = font_set.find(f"{DRAWINGML}{script}")
            if element is not None:
                element.set("typeface", FONT)
    content[theme_name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with tempfile.NamedTemporaryFile(suffix=".docx", dir=path.parent, delete=False) as handle:
        temporary = Path(handle.name)
    try:
        with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
            for name, data in content.items():
                target.writestr(name, data)
        temporary.replace(path)
    finally:
        if temporary.exists():
            temporary.unlink()

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None: node = OxmlElement(f"w:{side}"); tcMar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); elem = OxmlElement("w:tblHeader"); elem.set(qn("w:val"), "true"); trPr.append(elem)

def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PULSE Internal  |  Page "); set_font(run, 8, color=MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); paragraph._p.append(fld)

def configure(doc, short_title):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.85); sec.left_margin = sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)
    styles = doc.styles
    normal = styles["Normal"]; apply_style_font(normal); normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for key in ("Title", "Subtitle", "List Bullet", "List Number", "Caption"):
        apply_style_font(styles[key])
    for key, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, "1F4D78")):
        st = styles[key]; apply_style_font(st); st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color); st.font.bold = True
        st.paragraph_format.space_before = Pt(14); st.paragraph_format.space_after = Pt(6); st.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]; header.text = short_title; header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in header.runs: set_font(r, 8, bold=True, color=MUTED)
    page_number(sec.footer.paragraphs[0])

def title_page(doc, title, subtitle, audience, hero=None):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(34); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists(): p.add_run().add_picture(str(LOGO), width=Inches(2.0))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title); set_font(r, 26, bold=True, color="163B70")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(18)
    r = p.add_run(subtitle); set_font(r, 12.5, color=MUTED)
    table = doc.add_table(rows=2, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    values = [("Audience", audience), ("Release", "2026.07.22"), ("Classification", "Internal"), ("Review", "After a material system change")]
    for i, (label, val) in enumerate(values):
        cell = table.cell(i // 2, i % 2); cell.width = Inches(3.2); shade(cell, PALE); set_cell_margins(cell)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n"); set_font(r, 8, bold=True, color=MUTED)
        r = p.add_run(val); set_font(r, 9.5, bold=True, color=INK)
    if hero and hero.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(22)
        p.add_run().add_picture(str(hero), width=Inches(6.35))
    doc.add_page_break()

def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); set_font(r, 10.5, bold=True)
        r = p.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    return p

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item); set_font(r)

def steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item); set_font(r)

def image(doc, file, caption):
    path = SHOT / file
    if path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(path), width=Inches(6.25))
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
        r = p.add_run(caption); set_font(r, 8.5, color=MUTED)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, "E8EEF5"); set_cell_margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths: c.width = Inches(widths[i])
        r = c.paragraphs[0].add_run(h); set_font(r, 9, bold=True, color="1F4D78")
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: cells[i].width = Inches(widths[i])
            r = cells[i].paragraphs[0].add_run(str(val)); set_font(r, 8.8)
    doc.add_paragraph()

def user_guide():
    d = Document(); configure(d, "PULSE User Guide")
    title_page(d, "PULSE User Guide", "A practical guide to the current work, review, meeting, and travel tools.", "All authorized users", SHOT / "01-dashboard-current.png")
    d.add_heading("1. Start here", 1); para(d, "Open PULSE from the approved SharePoint page. Your SharePoint/CAC session identifies you automatically; PULSE does not use a separate password.")
    bullets(d, ["Use Dashboard to start a work session and find the major tools.", "Use the bell for attention items and the user menu for notification preferences.", "Outside SharePoint, PULSE runs in non-persistent local fallback mode. Do not treat fallback data as an official record."])
    d.add_heading("2. Use the dashboard and overview", 1); image(d, "02-overview-current.png", "Current Team Overview — portfolio health, workload, resources, and operational queues.")
    para(d, "Dashboard is the start-of-session control point. Use Overview → My for your tasks and assigned projects, or Overview → Team for the live portfolio summary. Make operational changes in the underlying project, travel, or review workspace; overview is a decision-support view.")
    d.add_heading("3. Manage project work", 1); image(d, "03-project-tracker-current.png", "Current project Tracker — owners, dates, health, risks, and reporting are reachable from one workspace.")
    steps(d, ["Open Projects and select the required project, or click New Project to initiate a new effort.", "Confirm role assignments, lifecycle details, dates, and current health before editing.", "Create tracker records with a clear owner, practical due date, status, and execution context.", "Use Blocked only for work that cannot proceed without a dependency; name the dependency and raise it in the appropriate meeting."])
    d.add_heading("4. Run the meeting workflow", 1); image(d, "04-weekly-meeting-current.png", "Current Weekly Meeting workspace.")
    steps(d, ["Confirm the meeting scope and the participant roster.", "Review Rocks, priorities, decisions, and blocked actions.", "Capture discussion notes, then create or update named tracker actions before closing the session."])
    d.add_heading("5. Submit travel and route documents", 1); image(d, "05-travel-request-current.png", "Current Travel wizard — request type is the first of five guided steps.")
    steps(d, ["For travel, select Travel → New Request, complete the guided form, review it, and submit it for the required approval path.", "For formal concurrence, use Document Review with the correct revision, a due date, and the required reviewers—not only the project library.", "Reviewers record decisions and actionable feedback in PULSE. Keep the revision trail intact.", "After review completion, if signatures are required, upload the Final Pack PDF for sequential signing."])
    image(d, "06-document-review-current.png", "Current Document Review board — review and sign workflows are separated by purpose.")
    d.add_heading("6. Get help", 1); para(d, "For a defect, provide the time, SharePoint page/route, project or record identifier, your role, expected behavior, actual behavior, and an approved screenshot. Do not modify roles or SharePoint schema unless you are an authorized administrator.")
    save_doc(d, OUT / "PULSE-User-Guide.docx")

def technical_handoff():
    d = Document(); configure(d, "PULSE Technical Handoff")
    title_page(d, "PULSE Technical Handoff", "Architecture, SharePoint boundaries, packaging, validation, and support guidance.", "Developers, site administrators, and release managers", DIAGRAM / "pulse-executive-architecture-final.png")
    d.add_heading("1. System model", 1); para(d, "PULSE is a browser-only single-page application hosted by SharePoint / Firepit. It uses the user’s same-origin SharePoint session and REST endpoints; SharePoint Lists and libraries are the shared operational data source.")
    bullets(d, ["Keep raw SharePoint REST activity centralized in sharepoint-adapter.js.", "Preserve sharepoint-repo.js mapper behavior whenever changing list fields, especially *Json fields.", "Do not introduce an application server, required Azure registration, core Microsoft Graph dependency, or runtime CDN dependency into the ship artifact.", "Resolve business roles from PULSE App Roles; site administration alone is not the normal role model."])
    d.add_heading("2. Maintainable repository map", 1)
    table(d, ["Area", "Primary path", "Responsibility"], [("Shell", "index.html; assets/js/app.js", "Boot, routing, shared state, navigation"), ("Feature UI", "assets/js/pages/", "Module-specific rendering and interactions"), ("SharePoint", "sharepoint-adapter.js; sharepoint-repo.js", "Identity, REST, mapping, saves/removes"), ("Schema", "sharepoint-schema.js", "Lists, columns, validation, setup"), ("Ship build", "scripts/build-sharepoint-package.js", "Preferred single-file Firepit/SharePoint artifact")], [1.0, 2.5, 2.75])
    d.add_heading("3. Data and integration", 1)
    pic = DIAGRAM / "pulse-data-relationship-graph-final.png"
    if pic.exists():
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(pic), width=Inches(6.2))
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("PULSE SharePoint data relationships."); set_font(r, 8.5, color=MUTED)
    para(d, "PULSE loads the required lists into an in-memory store, uses a short-lived session cache, and refreshes safely around writes and modals. Saves are debounced and chained per object. List-item writes are last-write-wins, so data model changes require live read/write validation.")
    d.add_heading("4. Change and deployment workflow", 1)
    steps(d, ["Read the affected page module, mapper, and schema definition before editing.", "Confirm current-user REST, site detection, and list diagnostics in the intended environment.", "Test the smallest safe change locally, then in a safe SharePoint site.", "Build the release with scripts/build-sharepoint-package.js.", "Upload the artifact and validate on the actual Firepit-hosted page—not only in a local browser.", "Record the release identifier, validation evidence, and prior known-good package for rollback."])
    d.add_heading("5. First response troubleshooting", 1)
    table(d, ["Symptom", "First checks", "Safe response"], [("Blank app", "Firepit page, console, package artifact", "Rebuild the preferred one-file package; inspect asset paths/query strings."), ("Wrong role", "currentuser REST, site URL, PULSE App Roles", "Correct the active role record; do not assume site membership is enough."), ("Save fails", "Logs, permissions, list/field schema, mapper", "Validate the live request and mapping before changing page UI."), ("Wrong data", "*Json shape, mapper, cached/live record", "Correct mapper/schema together and test read/write."), ("Release defect", "Actual Firepit page and prior package", "Rollback first if validation fails, then investigate.")], [1.05, 2.35, 2.85])
    d.add_heading("6. Release acceptance", 1); bullets(d, ["Identity and PULSE role resolve correctly.", "Required list/column validation passes.", "Affected feature supports the required read/write paths.", "The generated HTML has no external runtime dependency or invalid asset query string.", "The real Firepit page opens with no console errors and a rollback package is retained."])
    save_doc(d, OUT / "PULSE-Technical-Handoff.docx")

def sop():
    d = Document(); configure(d, "PULSE Operations SOP")
    title_page(d, "PULSE Operations SOP", "Repeatable controls for daily work, projects, meetings, travel, reviews, administration, and releases.", "All users; approvers, administrators, and release managers as assigned", SHOT / "01-dashboard-current.png")
    procedures = [
        ("SOP 01 — Start-of-session control check", "Every user | At the start of each work session", ["Open the approved SharePoint-hosted PULSE page.", "Confirm your displayed name and role.", "Review Dashboard attention items, My Work, and red/amber projects.", "Open, delegate, or schedule each action requiring attention."], "Completion: active work is identified or explicitly delegated."),
        ("SOP 02 — Create and maintain project work", "Project lead or delegate | At creation and whenever scope, owners, dates, or health change", ["Navigate to Projects and select New Project. Enter Project name, Portfolios, and Description.", "Add Members using the people picker to establish the roster, then save.", "Open the project and configure Settings: set Lifecycle status, Technical health, and essential Dates (start, due, completion).", "Create tracker records with an owner, Start date, Due date, Action status, and Health.", "For blocked work, select Blocked status and document the dependency in the Blocked reason field.", "Review health in the next project meeting and update the plan."], "Control: red/amber conditions require a tracked cause, owner, and recovery action."),
        ("SOP 03 — Conduct weekly and project meetings", "Facilitator / project lead | Scheduled cadence", ["Confirm the meeting scope and roster.", "Review prior actions, Rocks, priorities, and blockers.", "Capture material decisions and notes.", "Convert agreed follow-ups into named actions with dates.", "Verify actions appear in the appropriate tracker before close."], "Control: meeting notes do not replace assigned tracker records."),
        ("SOP 04 — Submit and approve travel", "Traveler submits; authorized approver decides | Before authorization and after material changes", ["Complete the guided request with accurate purpose, dates, travelers, location, estimates, and support detail.", "Submit and monitor the request; update it if the plan changes.", "Approvers verify information, approve/return/deny with an explanatory note, and confirm the visible status.", "Complete the required debrief after travel."], "Control: calendar appearance is not authorization evidence."),
        ("SOP 05 — Route a document for formal review", "Document owner and assigned reviewers | When concurrence or auditable review is required", ["Navigate to Document Review and click New Review. Provide the Title, Project, Due Date, and Revision Label (e.g., v1.0).", "Upload the Current File representing the exact version to be reviewed.", "Add individual Reviewers or groups. Designate sequential signers if signatures are required.", "Reviewers receive notification and log their decision (Approved or Changes Requested) with actionable comments.", "If changes are requested, update the file, upload it as a new revision, and restart the review cycle. Do not overwrite history.", "When all reviewers approve, the record enters Review Complete. If signatures are required, upload the Final Pack PDF to begin sequential signing."], "Control: never replace an approved revision with unreviewed content."),
        ("SOP 06 — Administer configuration and roles", "Authorized PULSE administrator | Onboarding, offboarding, setup, or configuration change", ["Confirm current-user REST and target-site diagnostics.", "Manage active PULSE App Roles using the approved role model.", "Run SharePoint Setup to validate/provision schema; do not manually improvise list changes.", "Review activity records after sensitive changes and preserve at least one recovery-capable administrator."], "Control: site membership alone does not define normal PULSE permissions."),
        ("SOP 07 — Release and rollback", "Release manager / authorized maintainer | Every production change", ["Test locally and in a safe SharePoint environment.", "Build with build-sharepoint-package.js and retain the prior known-good package.", "Upload the release and validate in the real Firepit-hosted page.", "If validation fails, restore the prior package and capture failure details before further edits."], "Control: local browser validation does not replace Firepit validation.")
    ]
    for idx, (heading, owner, items, control) in enumerate(procedures):
        d.add_heading(heading, 1); para(d, owner, bold_lead=owner.split(" | ")[0]); steps(d, items); para(d, control, bold_lead="Control: ")
        if idx == 1: image(d, "03-project-tracker-current.png", "Current project Tracker.")
        if idx == 3: image(d, "05-travel-request-current.png", "Current guided Travel request form.")
        if idx == 4: image(d, "06-document-review-current.png", "Current Document Review workflow board.")
    d.add_heading("Incident reporting minimum", 1); para(d, "For any defect, record time, site/page, user role, route, affected record ID, expected result, actual result, browser/device, and an approved screenshot. Preserve Logs/audit evidence. Do not make uncoordinated production schema or role changes while troubleshooting.")
    save_doc(d, OUT / "PULSE-Operations-SOP.docx")

if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    user_guide(); technical_handoff(); sop()
    print("Built current PULSE DOCX set.")
